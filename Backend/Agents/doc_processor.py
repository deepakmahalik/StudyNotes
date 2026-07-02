"""
DOC / DOCX Processor
--------------------
Extracts text from Word documents and structures it into mind-map nodes
using the shared Gemini LLM chain from ocr_node_agent.
"""

import os
import json
import logging

logger = logging.getLogger(__name__)

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a .docx file preserving paragraph structure."""
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading level as indentation hint
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
            else:
                parts.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from legacy .doc files.
    Requires either antiword (CLI) or python-docx2txt as fallback.
    """
    # Try docx2txt first (handles some .doc files)
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            return text.strip()
    except Exception:
        pass

    # Try antiword via subprocess (Linux/WSL only)
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass

    raise ValueError(
        "Cannot read .doc file. Please convert it to .docx and try again, "
        "or install antiword / docx2txt."
    )


# ── Node generation ───────────────────────────────────────────────────────────

def process_doc(file_path: str, output_path: str) -> dict:
    """
    Extract text from *file_path* (.doc or .docx), send to Gemini, save
    structured nodes to *output_path*, and return the nodes dict.
    """
    ext = os.path.splitext(file_path)[1].lower()
    logger.info("DOC processor: extracting text from %s", file_path)

    if ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        text = extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported document format: {ext}")

    if not text or not text.strip():
        raise ValueError("No text could be extracted from the document.")

    logger.info("DOC processor: %d chars extracted, sending to LLM", len(text))

    from ocr_node_agent import chain
    response = chain.invoke({"ocr_text": text})
    raw = response.content.strip()

    # Strip markdown fences if present
    if raw.startswith("```json"):
        raw = raw[7:]
    if raw.startswith("```"):
        raw = raw[3:]
    if raw.endswith("```"):
        raw = raw[:-3]
    raw = raw.strip()

    nodes_data = json.loads(raw)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(nodes_data, f, indent=2)

    logger.info("DOC processor: nodes saved to %s", output_path)

    # Save copy to public folder for auto-import
    base_dir = os.path.dirname(os.path.abspath(__file__))
    public_ocr_path = os.path.abspath(os.path.join(base_dir, "..", "..", "MindForge", "public", "ocr_nodes.json"))
    try:
        os.makedirs(os.path.dirname(public_ocr_path), exist_ok=True)
        with open(public_ocr_path, "w", encoding="utf-8") as pub_f:
            json.dump(nodes_data, pub_f, indent=2)
        logger.info("DOC processor: copied nodes to MindForge public folder")
    except Exception as pe:
        logger.warning("DOC processor: could not copy to public folder: %s", pe)

    return nodes_data
